from pathlib import Path

import pytest

from app.ingestion.parsers import get_parser_for_path
from app.ingestion.parsers.base import ParserError


# ── TXT ───────────────────────────────────────────────────────────

def test_txt_parser_normalizes_utf8_text(tmp_path: Path) -> None:
    """TXT parser should normalize line endings and preserve line metadata."""

    path = tmp_path / "notes.txt"
    path.write_bytes("﻿first line\r\nsecond line\n\n".encode("utf-8"))

    sections = get_parser_for_path(path).parse(path)

    assert len(sections) == 1
    assert sections[0].text == "first line\nsecond line"
    assert sections[0].source_metadata == {"line_start": 1, "line_end": 2}


# ── Markdown ──────────────────────────────────────────────────────

def test_markdown_parser_preserves_markdown_text(tmp_path: Path) -> None:
    """Markdown parser should ingest readable Markdown as text."""

    path = tmp_path / "runbook.md"
    path.write_text("# Runbook\n\n- Check latency\n- Review citations\n", encoding="utf-8")

    sections = get_parser_for_path(path).parse(path)

    assert len(sections) == 1
    assert sections[0].text == "# Runbook\n- Check latency\n- Review citations"
    assert sections[0].source_metadata == {
        "format": "markdown",
        "line_start": 1,
        "line_end": 4,
    }


# ── DOCX ──────────────────────────────────────────────────────────

def test_docx_parser_returns_non_empty_paragraphs(tmp_path: Path) -> None:
    """DOCX parser should extract readable paragraph text."""

    from docx import Document as DocxDocument

    path = tmp_path / "brief.docx"
    document = DocxDocument()
    document.add_paragraph("Alpha paragraph")
    document.add_paragraph("")
    document.add_paragraph("Beta paragraph")
    document.save(path)

    sections = get_parser_for_path(path).parse(path)

    assert [section.text for section in sections] == [
        "Alpha paragraph",
        "Beta paragraph",
    ]
    assert sections[0].source_metadata == {
        "type": "paragraph",
        "paragraph_start": 1,
        "paragraph_end": 1,
        "block_index": 0,
    }


def test_docx_parser_preserves_paragraph_table_order(tmp_path: Path) -> None:
    """Paragraphs and tables should appear in document order, not grouped."""

    from docx import Document as DocxDocument

    path = tmp_path / "interleaved.docx"
    doc = DocxDocument()
    doc.add_paragraph("First paragraph")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "H1"
    table.cell(0, 1).text = "H2"
    table.cell(1, 0).text = "A1"
    table.cell(1, 1).text = "B1"

    doc.add_paragraph("After table paragraph")
    doc.save(path)

    sections = get_parser_for_path(path).parse(path)

    types_in_order = [s.source_metadata.get("type") for s in sections]

    assert types_in_order[0] == "paragraph"
    assert sections[0].text == "First paragraph"
    assert sections[0].source_metadata["block_index"] == 0

    assert "table" in types_in_order
    assert types_in_order[-1] == "paragraph"
    assert sections[-1].text == "After table paragraph"
    assert sections[-1].source_metadata["block_index"] == 2


def test_docx_parser_table_with_headers_and_data(tmp_path: Path) -> None:
    """Standard table with headers and multiple data rows."""

    from docx import Document as DocxDocument

    path = tmp_path / "standard_table.docx"
    doc = DocxDocument()
    table = doc.add_table(rows=4, cols=3)
    headers = ["Name", "Role", "Location"]
    data = [
        ["Alice", "Engineer", "NYC"],
        ["Bob", "Designer", "SF"],
        ["Carol", "Manager", "London"],
    ]
    for col, h in enumerate(headers):
        table.cell(0, col).text = h
    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, val in enumerate(row_data):
            table.cell(row_idx, col_idx).text = val
    doc.save(path)

    sections = get_parser_for_path(path).parse(path)

    types = [s.source_metadata.get("type") for s in sections]
    assert "table" in types
    assert "table_header" in types
    assert types.count("table_row") == 3

    # Table section should contain all data and enriched metadata
    table_section = next(s for s in sections if s.source_metadata["type"] == "table")
    for name in ["Name", "Alice", "Bob", "Carol"]:
        assert name in table_section.text
    assert table_section.source_metadata["format"] == "docx"
    assert table_section.source_metadata["column_count"] == 3
    assert table_section.source_metadata["row_count"] == 3
    assert table_section.source_metadata["headers"] == headers

    # Header section
    header_section = next(s for s in sections if s.source_metadata["type"] == "table_header")
    assert "Name" in header_section.text
    assert header_section.source_metadata["table_index"] == 0
    assert header_section.source_metadata["row"] == 0

    # Row sections should be labeled with column headers
    alice_section = next(
        s for s in sections
        if s.source_metadata["type"] == "table_row" and "Alice" in s.text
    )
    assert "Name: Alice" in alice_section.text or "Alice" in alice_section.text
    assert alice_section.source_metadata["table_index"] == 0


def test_docx_parser_merged_caption_via_real_merge(tmp_path: Path) -> None:
    """A merged caption row created with .merge() should be detected.

    Uses the real python-docx ``.merge()`` API to verify that the
    XML-level cell iteration avoids the double-expansion bug described
    in the progress document Section 2.4.
    """

    from docx import Document as DocxDocument

    path = tmp_path / "real_merge_caption.docx"
    doc = DocxDocument()
    table = doc.add_table(rows=3, cols=3)

    # Merge first row cells into a single caption
    table.cell(0, 0).merge(table.cell(0, 2))
    table.cell(0, 0).text = "Server Configuration Table"

    # Headers
    table.cell(1, 0).text = "Server"
    table.cell(1, 1).text = "IP"
    table.cell(1, 2).text = "Status"

    # Data
    table.cell(2, 0).text = "web-01"
    table.cell(2, 1).text = "10.0.0.1"
    table.cell(2, 2).text = "active"

    doc.save(path)
    sections = get_parser_for_path(path).parse(path)

    # The table section must show correct column count (3, not 9)
    table_section = next(s for s in sections if s.source_metadata["type"] == "table")
    assert table_section.source_metadata["column_count"] == 3

    # Caption text must appear in the table section
    assert "Server Configuration Table" in table_section.text

    types = [s.source_metadata.get("type") for s in sections]
    assert "table" in types
    assert "table_header" in types
    assert types.count("table_row") >= 1


def test_docx_parser_merged_caption_row(tmp_path: Path) -> None:
    """Legacy test: manual XML gridSpan on first row should still detect caption."""

    from docx import Document as DocxDocument
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    path = tmp_path / "merged_caption.docx"
    doc = DocxDocument()
    table = doc.add_table(rows=3, cols=3)

    caption_cell = table.cell(0, 0)
    caption_cell.text = "Server Configuration Table"
    tc_pr = caption_cell._tc.find(qn("w:tcPr"))
    if tc_pr is None:
        tc_pr = OxmlElement("w:tcPr")
        caption_cell._tc.insert(0, tc_pr)
    grid_span = OxmlElement("w:gridSpan")
    grid_span.set(qn("w:val"), "3")
    tc_pr.append(grid_span)

    table.cell(1, 0).text = "Server"
    table.cell(1, 1).text = "IP"
    table.cell(1, 2).text = "Status"
    table.cell(2, 0).text = "web-01"
    table.cell(2, 1).text = "10.0.0.1"
    table.cell(2, 2).text = "active"

    doc.save(path)
    sections = get_parser_for_path(path).parse(path)

    types = [s.source_metadata.get("type") for s in sections]
    assert "table" in types

    table_section = next(s for s in sections if s.source_metadata["type"] == "table")
    assert "Server Configuration Table" in table_section.text
    assert "table_header" in types
    assert types.count("table_row") >= 1


def test_docx_parser_preserves_empty_cells(tmp_path: Path) -> None:
    """Empty cells should be preserved as empty strings, not skipped.

    Column position must not shift — a value in column 3 must stay
    associated with column 3 even when column 2 is empty.
    """

    from docx import Document as DocxDocument

    path = tmp_path / "empty_cells.docx"
    doc = DocxDocument()
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "ColA"
    table.cell(0, 1).text = "ColB"
    table.cell(0, 2).text = "ColC"
    table.cell(1, 0).text = "val_a"
    table.cell(1, 2).text = "val_c"
    doc.save(path)

    sections = get_parser_for_path(path).parse(path)

    table_section = next(s for s in sections if s.source_metadata["type"] == "table")
    assert "val_c" in table_section.text

    row_section = next(
        s for s in sections
        if s.source_metadata["type"] == "table_row" and "val_a" in s.text
    )
    pos_a = row_section.text.index("val_a")
    pos_c = row_section.text.index("val_c")
    assert pos_a < pos_c


def test_docx_parser_produces_table_level_section(tmp_path: Path) -> None:
    """Every table should produce a table-level section with complete text."""

    from docx import Document as DocxDocument

    path = tmp_path / "two_tables.docx"
    doc = DocxDocument()

    for t_idx in range(2):
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = f"T{t_idx}A"
        table.cell(0, 1).text = f"T{t_idx}B"
        table.cell(1, 0).text = f"T{t_idx}C"
        table.cell(1, 1).text = f"T{t_idx}D"

    doc.save(path)
    sections = get_parser_for_path(path).parse(path)

    table_sections = [s for s in sections if s.source_metadata["type"] == "table"]
    assert len(table_sections) == 2

    assert "T0A" in table_sections[0].text
    assert "T0D" in table_sections[0].text
    assert "T1A" in table_sections[1].text
    assert "T1D" in table_sections[1].text

    assert table_sections[0].source_metadata["table_index"] == 0
    assert table_sections[1].source_metadata["table_index"] == 1
    # Distinct block_index values
    assert table_sections[0].source_metadata["block_index"] != table_sections[1].source_metadata["block_index"]


def test_docx_parser_empty_table_produces_no_sections(tmp_path: Path) -> None:
    """A completely empty table (no rows) should not add sections."""

    from docx import Document as DocxDocument

    path = tmp_path / "empty_table.docx"
    doc = DocxDocument()
    doc.add_table(rows=0, cols=0)
    doc.save(path)

    with pytest.raises(ParserError, match="no usable text"):
        get_parser_for_path(path).parse(path)


def test_docx_parser_table_without_headers(tmp_path: Path) -> None:
    """Table with data that does not look like headers should skip header detection.

    Uses long values that exceed the 50-char header heuristic threshold so
    both rows are treated as data.
    """

    from docx import Document as DocxDocument

    path = tmp_path / "no_header_table.docx"
    doc = DocxDocument()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Data center primary network infrastructure rack 42A"
    table.cell(0, 1).text = "This server has been operational since January 2024"
    table.cell(1, 0).text = "Backup node secondary network infrastructure rack 17C"
    table.cell(1, 1).text = "Decommissioned after three years of continuous service"
    doc.save(path)

    sections = get_parser_for_path(path).parse(path)

    types = [s.source_metadata.get("type") for s in sections]
    assert "table" in types
    assert "table_header" not in types
    assert types.count("table_row") == 2


def test_docx_parser_low_confidence_header_preserves_first_row(tmp_path: Path) -> None:
    """When header confidence is below threshold, the first row is kept as data.

    Short value pairs like IP + status look like headers to the heuristic
    but the vocabulary-overlap and data-pattern checks both fail, resulting
    in low confidence.  The parser must NOT drop the first row.
    """

    from docx import Document as DocxDocument

    path = tmp_path / "short_data.docx"
    doc = DocxDocument()
    table = doc.add_table(rows=3, cols=2)
    # These values are short but are data, not headers
    table.cell(0, 0).text = "web-01"
    table.cell(0, 1).text = "active"
    table.cell(1, 0).text = "web-02"
    table.cell(1, 1).text = "standby"
    table.cell(2, 0).text = "web-03"
    table.cell(2, 1).text = "inactive"
    doc.save(path)

    sections = get_parser_for_path(path).parse(path)

    types = [s.source_metadata.get("type") for s in sections]
    assert "table" in types

    # Low confidence: first row must NOT be stripped
    # All 3 rows should appear as table_row sections
    assert types.count("table_row") == 3

    # The first row text should be among the row sections
    row_texts = [s.text for s in sections if s.source_metadata.get("type") == "table_row"]
    assert any("web-01" in t for t in row_texts), "first data row must not be lost"


def test_docx_parser_tbl_header_marker_detected(tmp_path: Path) -> None:
    """Word's w:tblHeader repeat-header marker should give high confidence."""

    from docx import Document as DocxDocument
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    path = tmp_path / "tbl_header_marker.docx"
    doc = DocxDocument()
    table = doc.add_table(rows=3, cols=2)

    # Add w:tblHeader marker to the first row
    tr_pr = OxmlElement("w:trPr")
    tbl_header = OxmlElement("w:tblHeader")
    tr_pr.append(tbl_header)
    table.rows[0]._tr.insert(0, tr_pr)

    table.cell(0, 0).text = "Server"
    table.cell(0, 1).text = "Status"
    table.cell(1, 0).text = "web-01"
    table.cell(1, 1).text = "active"
    table.cell(2, 0).text = "web-02"
    table.cell(2, 1).text = "standby"
    doc.save(path)

    sections = get_parser_for_path(path).parse(path)

    types = [s.source_metadata.get("type") for s in sections]

    # With w:tblHeader, confidence should be 0.9 — headers detected
    assert "table_header" in types
    # Data rows: only 2 rows (first row is header)
    assert types.count("table_row") == 2

    table_section = next(s for s in sections if s.source_metadata["type"] == "table")
    assert table_section.source_metadata["header_confidence"] == 0.9


# ── XLSX ──────────────────────────────────────────────────────────

def test_xlsx_parser_returns_visible_sheet_rows(tmp_path: Path) -> None:
    """XLSX parser should convert non-empty rows into deterministic text."""

    from openpyxl import Workbook

    path = tmp_path / "table.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Data"
    sheet.append(["Name", "Value"])
    sheet.append(["Alpha", 42])
    workbook.save(path)

    sections = get_parser_for_path(path).parse(path)

    assert [section.text for section in sections] == ["Name | Value", "Alpha | 42"]
    assert sections[1].source_metadata == {
        "sheet_name": "Data",
        "row_start": 2,
        "row_end": 2,
    }


# ── PDF ───────────────────────────────────────────────────────────

def test_pdf_parser_returns_one_section_per_text_page(tmp_path: Path) -> None:
    """PDF parser should extract text page by page."""

    import fitz

    path = tmp_path / "paper.pdf"
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "PDF page text")
    pdf.save(path)
    pdf.close()

    sections = get_parser_for_path(path).parse(path)

    assert len(sections) == 1
    assert "PDF page text" in sections[0].text
    assert sections[0].source_metadata == {"page_number": 1}


# ── Common ────────────────────────────────────────────────────────

def test_empty_document_raises_parser_error(tmp_path: Path) -> None:
    """Parsers should fail clearly when no usable text exists."""

    path = tmp_path / "empty.txt"
    path.write_text("", encoding="utf-8")

    with pytest.raises(ParserError, match="no usable text"):
        get_parser_for_path(path).parse(path)
