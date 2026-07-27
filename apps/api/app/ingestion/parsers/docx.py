from pathlib import Path

from app.ingestion.parsers.base import NormalizedSection, ParserError, StructuredTable


def _looks_like_header(cells: list[str]) -> bool:
    """Return True when cells look like column labels rather than data.

    Heuristic: header cells are short, do not end with sentence punctuation,
    and at least one cell is non-empty.
    """
    non_empty = [c for c in cells if c]
    if not non_empty:
        return False
    for cell in non_empty:
        if len(cell) > 50:
            return False
        if cell.endswith((".", "。", "!", "？", "?", "；", ";", "：", ":")):
            return False
    return True


def _label_row(values: list[str], headers: list[str]) -> list[str]:
    """Pair each non-empty cell value with its column header.

    Returns strings like "ServerName: web-01".
    Cells without a matching header keep their raw value.
    """
    labeled: list[str] = []
    for col_index, value in enumerate(values):
        if not value:
            continue
        if col_index < len(headers) and headers[col_index]:
            labeled.append(f"{headers[col_index]}: {value}")
        else:
            labeled.append(value)
    return labeled


def _iter_body_elements(document):
    """Yield (tag, index, block_index) tuples in document order."""
    from docx.oxml.ns import qn

    tag_p = qn("w:p")
    tag_tbl = qn("w:tbl")

    para_index = 0
    table_index = 0
    block_index = 0

    for child in document.element.body:
        if child.tag == tag_p:
            yield ("p", para_index, block_index)
            para_index += 1
            block_index += 1
        elif child.tag == tag_tbl:
            yield ("tbl", table_index, block_index)
            table_index += 1
            block_index += 1


def _extract_paragraph_text(para_index: int, document) -> str | None:
    """Extract stripped text from a paragraph at the given index."""
    paragraphs = document.paragraphs
    if para_index >= len(paragraphs):
        return None
    text = paragraphs[para_index].text.strip()
    return text if text else None


def _extract_cell_text(tc_element) -> tuple[str, bool]:
    """Extract text from a ``w:tc`` XML element.

    Returns ``(text, has_nested_table)``.  Nested-table text is included
    for fallback discoverability but flagged so callers can distinguish
    it from intentional parent-cell content.
    """
    from docx.oxml.ns import qn

    w_t = qn("w:t")
    w_tbl = qn("w:tbl")

    parts: list[str] = []
    nested_parts: list[str] = []
    has_nested = False

    for t_node in tc_element.iter(w_t):
        # Check whether this text run is inside a nested w:tbl
        ancestor = t_node.getparent()
        in_nested = False
        while ancestor is not None and ancestor != tc_element:
            if ancestor.tag == w_tbl:
                in_nested = True
                has_nested = True
                break
            ancestor = ancestor.getparent()
        if in_nested:
            if t_node.text:
                nested_parts.append(t_node.text)
        else:
            if t_node.text:
                parts.append(t_node.text)

    text = " ".join(parts).strip()
    if nested_parts:
        nested_text = " ".join(nested_parts).strip()
        if nested_text:
            if text:
                text = f"{text} [nested table: {nested_text}]"
            else:
                text = f"[nested table: {nested_text}]"
    return text, has_nested


def _has_tbl_header(row_element) -> bool:
    """Return True when the ``w:tr`` element carries the Word repeated-header marker."""
    from docx.oxml.ns import qn

    tr_pr = row_element.find(qn("w:trPr"))
    if tr_pr is not None:
        return tr_pr.find(qn("w:tblHeader")) is not None
    return False


def _detect_headers_and_confidence(
    rows: list[list[str]],
    row_elements: list,
) -> tuple[list[str], float, str]:
    """Return (headers, confidence, source) for a table.

    Confidence levels:
    - 0.90  : ``w:tblHeader`` marker present → source="word_marker"
    - 0.75  : vocabulary non-overlap — header cells don't appear in data
    - 0.60  : data-pattern signal — data rows look different from header row
    - 0.40  : weak heuristic — first row looks header-like but can't confirm
    - 0.0   : no header detected → source="none"

    Data-quantity caps are NOT applied — a table with only 1-2 data rows
    is still analysed by comparing the first row against data columns.
    """
    if not rows or not row_elements:
        return [], 0.0, "none"

    first_row = rows[0]

    # Strong signal: Word repeated-header marker
    if _has_tbl_header(row_elements[0]):
        return first_row, 0.9, "word_marker"

    # Heuristic: do the cells look like labels?
    if not _looks_like_header(first_row):
        return [], 0.0, "none"

    if len(rows) < 2:
        return first_row, 0.4, "heuristic_weak"

    data_rows = rows[1:]
    num_cols = max((len(r) for r in rows), default=0)

    # ── Vocabulary non-overlap ──
    header_values_lower = {c.lower() for c in first_row if c}
    data_values_lower: set[str] = set()
    for row in data_rows:
        for cell in row:
            if cell:
                data_values_lower.add(cell.lower())
    vocab_distinct = bool(header_values_lower and not header_values_lower & data_values_lower)

    if vocab_distinct:
        confidence = _header_confidence_from_vocab(rows, num_cols)
        return first_row, confidence, "heuristic_vocab"

    # ── Data-pattern signal ──
    data_cell_count = 0
    label_like_count = 0
    for row in data_rows:
        for cell in row:
            if not cell:
                continue
            if len(cell) > 50 or any(c.isdigit() for c in cell):
                data_cell_count += 1
            elif len(cell) < 20 and not cell.endswith((".", "。", "!", "?")):
                label_like_count += 1

    if data_cell_count > label_like_count:
        return first_row, 0.6, "heuristic_pattern"

    return first_row, 0.4, "heuristic_weak"


def _header_confidence_from_vocab(rows: list[list[str]], num_cols: int) -> float:
    """Modulate vocabulary-non-overlap confidence based on column structure.

    When columns show strong data-like repeating patterns (common prefix
    like ``web-01/web-02/web-03`` or consistent non-alpha character
    classes), the first row is more likely unique data than a header.

    Returns 0.75 for clear headers, 0.55 for likely data.
    """
    if num_cols == 0:
        return 0.75

    structured_cols = 0
    for col_idx in range(num_cols):
        cells = [row[col_idx] for row in rows if col_idx < len(row) and row[col_idx]]
        if len(cells) < 2:
            continue

        # Check common prefix (3+ chars shared by ≥60% of cells)
        prefix = cells[0][:3] if len(cells[0]) >= 3 else ""
        if prefix and len(prefix) == 3:
            prefix_count = sum(1 for c in cells if c.startswith(prefix))
            if prefix_count >= len(cells) * 0.6:
                structured_cols += 1
                continue

        # Check consistent non-alpha character pattern
        from collections import Counter

        patterns = []
        for c in cells:
            stripped = c.replace(".", "").replace("-", "").replace("_", "").replace(" ", "")
            if not stripped:
                patterns.append("alpha")
            elif stripped.isdigit():
                patterns.append("numeric")
            elif any(ch.isdigit() for ch in stripped):
                patterns.append("alphanumeric")
            else:
                patterns.append("alpha")
        top_pattern, top_count = Counter(patterns).most_common(1)[0]
        if top_pattern != "alpha" and top_count >= len(cells) * 0.8:
            structured_cols += 1

    # Most columns are structured → likely data, not headers
    if num_cols == 1:
        return 0.55 if structured_cols >= 1 else 0.75
    if structured_cols * 2 >= num_cols:
        return 0.55
    return 0.75


def _normalize_table(table, table_index: int, block_index: int) -> StructuredTable:
    """Convert a python-docx Table into a StructuredTable.

    Uses ``iterchildren()`` (direct children only) to avoid pulling
    nested-table rows into the parent table.
    """
    from docx.oxml.ns import qn

    # ── Phase 1: iterate direct-child rows and cells ──
    raw_rows: list[list[tuple[str, int, str | None]]] = []
    row_elements: list = []
    max_cols = 0

    tbl_element = table._tbl
    w_tr = qn("w:tr")
    w_tc = qn("w:tc")

    nested_rows: set[int] = set()  # 0-based row indices with nested tables
    for tr in tbl_element.iterchildren(w_tr):
        row_index = len(row_elements)
        row_elements.append(tr)
        cells: list[tuple[str, int, str | None]] = []
        row_cols = 0
        row_has_nested = False
        for tc in tr.iterchildren(w_tc):
            text, has_nested = _extract_cell_text(tc)
            if has_nested:
                row_has_nested = True
            tc_pr = tc.find(qn("w:tcPr"))
            span = 1
            v_merge: str | None = None
            if tc_pr is not None:
                grid_span_el = tc_pr.find(qn("w:gridSpan"))
                if grid_span_el is not None:
                    span = int(grid_span_el.get(qn("w:val")))
                v_merge_el = tc_pr.find(qn("w:vMerge"))
                if v_merge_el is not None:
                    val = v_merge_el.get(qn("w:val"))
                    v_merge = val if val else "continue"
            cells.append((text, span, v_merge))
            row_cols += span
        if row_has_nested:
            nested_rows.add(row_index)
        max_cols = max(max_cols, row_cols)
        raw_rows.append(cells)

    if not raw_rows:
        return StructuredTable(
            table_index=table_index,
            block_index=block_index,
            source_metadata={"format": "docx", "table_index": table_index},
        )

    # ── Phase 2: expand gridSpan and handle vMerge ──
    normalized: list[list[str]] = []
    prev_expanded: list[str] = []

    for row_cells in raw_rows:
        expanded: list[str] = []
        new_prev: list[str] = []
        col = 0

        for text, span, v_merge in row_cells:
            if v_merge == "continue":
                for _ in range(span):
                    val = prev_expanded[col] if col < len(prev_expanded) else text
                    expanded.append(val)
                    new_prev.append(val)
                    col += 1
            else:
                for _ in range(span):
                    expanded.append(text)
                    col += 1
                new_prev.extend([text] * span)

        while len(expanded) < max_cols:
            expanded.append("")
        while len(new_prev) < max_cols:
            new_prev.append("")

        normalized.append(expanded)
        prev_expanded = new_prev

    # ── Phase 3: detect caption, headers, data rows ──
    caption: str | None = None
    headers: list[str] = []
    header_confidence: float = 0.0
    header_source: str = "none"
    data_start = 0

    # Detect caption via gridSpan
    first_raw = raw_rows[0]
    for text, span, _ in first_raw:
        if text and span > max_cols / 2:
            caption = text
            data_start = 1
            break

    # Fallback caption detection
    if caption is None:
        first_expanded = normalized[0]
        non_empty = [c for c in first_expanded if c]
        if len(non_empty) == 1 and max_cols > 1 and len(normalized) > 1:
            other_avg = (
                sum(sum(1 for c in row if c) for row in normalized[1:])
                / len(normalized[1:])
            )
            if other_avg > 1:
                caption = non_empty[0]
                data_start = 1

    # Detect header row with confidence
    if data_start < len(normalized):
        headers, header_confidence, header_source = _detect_headers_and_confidence(
            normalized[data_start:],
            row_elements[data_start:],
        )
        if headers and header_confidence >= 0.6:
            data_start += 1
        elif headers and header_confidence < 0.6:
            # Cannot confirm headers — keep first row as data,
            # use generated column names for embedding.
            headers = [f"Column {i + 1}" for i in range(max_cols)]
            header_confidence = 0.0
            header_source = "generated"
            # data_start unchanged (first row kept as data)

    # Map raw (0-based) nested row indices to data row indices.
    # caption/header rows are excluded from data rows.
    data_nested: list[int] = []
    for raw_idx in sorted(nested_rows):
        if raw_idx < data_start:
            continue  # nested table in caption or header — skip
        data_nested.append(raw_idx - data_start)

    return StructuredTable(
        table_index=table_index,
        block_index=block_index,
        caption=caption,
        headers=headers,
        rows=normalized[data_start:],
        column_count=max_cols,
        header_confidence=header_confidence,
        source_metadata={
            "format": "docx",
            "table_index": table_index,
            "header_source": header_source,
            "nested_table_fallback": len(data_nested) > 0,
            "nested_rows": data_nested,
        },
    )


def _emit_table_sections(
    table: StructuredTable,
    sections: list[NormalizedSection],
) -> None:
    """Append NormalizedSections for a StructuredTable to the sections list."""
    table_index = table.table_index
    block_index = table.block_index
    caption = table.caption
    headers = table.headers
    rows = table.rows
    column_count = table.column_count
    header_source = table.source_metadata.get("header_source", "none")
    has_nested = table.source_metadata.get("nested_table_fallback", False)
    nested_rows = table.source_metadata.get("nested_rows", [])

    if not rows and not headers:
        return

    # ── Complete-table section ──
    parts: list[str] = []
    if caption:
        parts.append(caption)
    if headers:
        parts.append(" | ".join(h for h in headers if h))
    for row in rows:
        parts.append(" | ".join(row))

    full_text = "\n".join(parts).strip()
    if full_text:
        sections.append(
            NormalizedSection(
                section_index=len(sections),
                text=full_text,
                source_metadata={
                    "format": "docx",
                    "type": "table",
                    "table_chunk_type": "table",
                    "table_index": table_index,
                    "block_index": block_index,
                    "caption": caption,
                    "headers": headers,
                    "column_count": column_count,
                    "row_count": len(rows),
                    "header_confidence": table.header_confidence,
                    "header_source": header_source,
                    "nested_table_fallback": has_nested,
                    "nested_rows": nested_rows if has_nested else [],
                },
            )
        )

    # ── Header section ──
    if headers:
        header_text = " | ".join(h for h in headers if h)
        if header_text:
            sections.append(
                NormalizedSection(
                    section_index=len(sections),
                    text=header_text,
                    source_metadata={
                        "format": "docx",
                        "type": "table_header",
                        "table_chunk_type": "table_header",
                        "table_index": table_index,
                        "block_index": block_index,
                        "row": 0,
                        "headers": headers,
                        "header_source": header_source,
                    },
                )
            )

    # ── Data row sections ──
    for row_index, row in enumerate(rows):
        if headers:
            labeled = _label_row(row, headers)
            row_text = " | ".join(labeled)
        else:
            row_text = " | ".join(row)
        if not row_text.strip(" |"):
            continue
        row_meta: dict = {
            "format": "docx",
            "type": "table_row",
            "table_chunk_type": "table_row",
            "table_index": table_index,
            "block_index": block_index,
            "data_row": row_index + 1,  # 1-based
            "headers": headers if headers else None,
        }
        if has_nested and row_index in nested_rows:
            row_meta["nested_table_fallback"] = True
        sections.append(
            NormalizedSection(
                section_index=len(sections),
                text=row_text,
                source_metadata=row_meta,
            )
        )


class DocxParser:
    """Parser for DOCX paragraph and table text."""

    def parse(self, path: Path) -> list[NormalizedSection]:
        from docx import Document

        document = Document(path)
        sections: list[NormalizedSection] = []

        for tag, index, block_index in _iter_body_elements(document):
            if tag == "p":
                text = _extract_paragraph_text(index, document)
                if text is None:
                    continue
                sections.append(
                    NormalizedSection(
                        section_index=len(sections),
                        text=text,
                        source_metadata={
                            "type": "paragraph",
                            "paragraph_start": index + 1,
                            "paragraph_end": index + 1,
                            "block_index": block_index,
                        },
                    )
                )
            elif tag == "tbl":
                if index >= len(document.tables):
                    continue
                table = _normalize_table(document.tables[index], index, block_index)
                _emit_table_sections(table, sections)

        if not sections:
            raise ParserError("Document contains no usable text")
        return sections
